from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
import os
from datetime import datetime
from core.logger import logger
from PyQt5.QtGui import QTextDocument

def export_to_docx(date, data, file_path, logo_path=None):
    """Esporta i dati in un file Word"""
    try:
        doc = Document()
        
        # Stile del documento
        sections = doc.sections
        for section in sections:
            section.page_height = Inches(11.69)  # A4
            section.page_width = Inches(8.27)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
        
        # Crea una tabella invisibile per il layout (1 riga, 2 colonne)
        layout_table = doc.add_table(rows=1, cols=2)
        layout_table.allow_autofit = True
        
        # Rimuovi i bordi dalla tabella di layout
        for row in layout_table.rows:
            for cell in row.cells:
                cell._tc.get_or_add_tcPr().append(parse_xml('<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>'))
        
        # Cella sinistra per il logo (solo se il logo esiste)
        logo_cell = layout_table.rows[0].cells[0]
        if logo_path and os.path.exists(logo_path):
            logo_paragraph = logo_cell.paragraphs[0]
            logo_run = logo_paragraph.add_run()
            logo_run.add_picture(logo_path, width=Inches(1.5))
        
        # Cella destra per il titolo
        title_cell = layout_table.rows[0].cells[1]
        title_paragraph = title_cell.paragraphs[0]
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run('Prenotazioni')
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = RGBColor(255, 0, 0)  # Rosso
        title_run.font.bold = True
        
        # Data
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_paragraph.add_run(f"Data: {date}")
        date_run.font.size = Pt(12)
        date_run.font.bold = True
        
        # Spazio prima della tabella
        doc.add_paragraph()
        
        # Tabella
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.allow_autofit = True
        
        # Stile intestazioni
        header_cells = table.rows[0].cells
        headers = ['Orario', 'Nome', 'Cognome', 'Prima Donazione']
        for i, text in enumerate(headers):
            cell = header_cells[i]
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(text)
            run.font.bold = True
            run.font.size = Pt(11)
            cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="E0E0E0"/>'))
        
        # Dati
        for orario, nome, cognome, prima_donazione, _ in data:
            row_cells = table.add_row().cells
            # Mostra "Sì" solo se è prima donazione, altrimenti stringa vuota
            prima_donazione_text = "Sì" if prima_donazione == "Sì" else ""
            for i, text in enumerate([orario, nome, cognome, prima_donazione_text]):
                cell = row_cells[i]
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run(str(text))
                run.font.size = Pt(11)
        
        # Imposta larghezze colonne
        widths = (Inches(1), Inches(2), Inches(2), Inches(1.5))
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_run = footer.add_run(f'Creato da Hemodos il {datetime.now().strftime("%d/%m/%Y alle %H:%M:%S")}')
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)  # Grigio
        
        # Salva nel percorso specificato
        doc.save(file_path)
        return True
        
    except Exception as e:
        logger.error(f"Errore nell'esportazione in Word: {str(e)}")
        return False

def print_data(printer, date, data, file_path, logo_path=None):
    """Crea il documento Word per la stampa"""
    try:
        doc = Document()
        
        # Stile del documento
        sections = doc.sections
        for section in sections:
            section.page_height = Inches(11.69)  # A4
            section.page_width = Inches(8.27)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
        
        # Crea una tabella invisibile per il layout (1 riga, 2 colonne)
        layout_table = doc.add_table(rows=1, cols=2)
        layout_table.allow_autofit = True
        
        # Rimuovi i bordi dalla tabella di layout
        for row in layout_table.rows:
            for cell in row.cells:
                cell._tc.get_or_add_tcPr().append(parse_xml('<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>'))
        
        # Cella sinistra per il logo (solo se il logo esiste)
        logo_cell = layout_table.rows[0].cells[0]
        if logo_path and os.path.exists(logo_path):
            logo_paragraph = logo_cell.paragraphs[0]
            logo_run = logo_paragraph.add_run()
            logo_run.add_picture(logo_path, width=Inches(1.5))
        
        # Cella destra per il titolo
        title_cell = layout_table.rows[0].cells[1]
        title_paragraph = title_cell.paragraphs[0]
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run('Prenotazioni')
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = RGBColor(255, 0, 0)  # Rosso
        title_run.font.bold = True
        
        # Data
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_paragraph.add_run(f"Data: {date}")
        date_run.font.size = Pt(12)
        date_run.font.bold = True
        
        # Spazio prima della tabella
        doc.add_paragraph()
        
        # Tabella
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.allow_autofit = True
        
        # Stile intestazioni
        header_cells = table.rows[0].cells
        headers = ['Orario', 'Nome', 'Cognome', 'Prima Donazione']
        for i, text in enumerate(headers):
            cell = header_cells[i]
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(text)
            run.font.bold = True
            run.font.size = Pt(11)
            cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="E0E0E0"/>'))
        
        # Dati
        for orario, nome, cognome, prima_donazione, _ in data:
            row_cells = table.add_row().cells
            # Mostra "Sì" solo se è prima donazione, altrimenti stringa vuota
            prima_donazione_text = "Sì" if prima_donazione == "Sì" else ""
            for i, text in enumerate([orario, nome, cognome, prima_donazione_text]):
                cell = row_cells[i]
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run(str(text))
                run.font.size = Pt(11)
        
        # Imposta larghezze colonne
        widths = (Inches(1), Inches(2), Inches(2), Inches(1.5))
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_run = footer.add_run(f'Creato da Hemodos il {datetime.now().strftime("%d/%m/%Y alle %H:%M:%S")}')
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)  # Grigio
        
        # Salva il documento
        doc.save(file_path)
        return True
        
    except Exception as e:
        logger.error(f"Errore nella creazione del documento: {str(e)}")
        return False
